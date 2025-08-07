#!/usr/bin/env python3

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, Inches
from dotenv import load_dotenv
from docx.oxml.shared import OxmlElement, qn
from docx2pdf import convert
import os
load_dotenv()

def set_paragraph_bottom_border(paragraph):
    """
    Adds a full-width bottom border (horizontal line) under a paragraph
    by editing its XML.
    """
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    for element in pBdr.findall(qn('w:bottom')):
        pBdr.remove(element)
    pBdr.append(bottom)


class McKinseyCVGenerator:
    """
    A class that encapsulates all logic to generate a McKinsey-style CV document,
    using data from a provided config dictionary.
    """

    def __init__(self, config=None, extra_skills=None, output_filename="Udit_Resume.docx"):
        """
        :param config: Dictionary loaded from YAML (contains personal_info, certifications, skills, etc.)
        :param extra_skills: Additional skills discovered from OpenAI (to merge into the 'Skills' section)
        :param output_filename: Where to save the final doc
        """
        if config is None:
            raise ValueError("No config dictionary provided to McKinseyCVGenerator!")

        self.config = config
        self.extra_skills = extra_skills or []
        self.output_filename = output_filename

        self.doc = docx.Document()
        self._set_margins()
        self._define_styles()

    def _set_margins(self):
        for section in self.doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

    def _define_styles(self):
        styles = self.doc.styles

        # Major heading style
        if 'CustomHeading' not in styles:
            heading_style = styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
            heading_style.font.name = 'Calibri'
            heading_style.font.size = Pt(12)
            heading_style.font.bold = True
            heading_style.paragraph_format.space_before = Pt(0)
            heading_style.paragraph_format.space_after = Pt(0)
            heading_style.paragraph_format.line_spacing = 1

        # Subheading style
        if 'CustomSubheading' not in styles:
            subheading_style = styles.add_style('CustomSubheading', WD_STYLE_TYPE.PARAGRAPH)
            subheading_style.font.name = 'Calibri'
            subheading_style.font.size = Pt(10)
            subheading_style.font.bold = True
            subheading_style.paragraph_format.space_before = Pt(0)
            subheading_style.paragraph_format.space_after = Pt(0)
            subheading_style.paragraph_format.line_spacing = 1

        # Normal style
        normal_style = styles['Normal']
        normal_style.font.name = 'Calibri'
        normal_style.font.size = Pt(10)
        normal_style.paragraph_format.space_before = Pt(0)
        normal_style.paragraph_format.space_after = Pt(0)
        normal_style.paragraph_format.line_spacing = 1

        # List Bullet style
        if 'List Bullet' in styles:
            bullet_style = styles['List Bullet']
            bullet_style.font.name = 'Calibri'
            bullet_style.font.size = Pt(9.5)
            bullet_style.paragraph_format.left_indent = Inches(0.4)
            bullet_style.paragraph_format.hanging_indent = Inches(0)
            bullet_style.paragraph_format.space_before = Pt(0)
            bullet_style.paragraph_format.space_after = Pt(0)
            bullet_style.paragraph_format.line_spacing = 1

    def _add_heading_with_line(self, text):
        p = self.doc.add_paragraph(style='CustomHeading')
        p.add_run(text)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_bottom_border(p)
        return p

    def add_bullet(self, text):
        p = self.doc.add_paragraph(style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p_format = p.paragraph_format
        p_format.left_indent = Inches(0.3)
        p_format.first_line_indent = Inches(-0.15)
        p.add_run(text)
        return p

    def _add_bold_label_value(self, label, value):
        p = self.doc.add_paragraph(style='Normal')
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run(label).bold = True
        p.add_run(value)
        return p

    def build(self):
        """
        Build the CV from the self.config data.
        """

        # ---- 1. Personal Info ----
        pi = self.config.get('personal_info', {})
        name_line = self.doc.add_paragraph(pi.get("name", ""), style='CustomHeading')
        name_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        phone_line = self.doc.add_paragraph(pi.get("phone", ""), style='Normal')
        phone_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        contact_line = self.doc.add_paragraph(f"Email: {pi.get('email','')}", style='Normal')
        contact_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Optional fields
        linkedin = pi.get("linkedin")
        if linkedin:
            linkedin_line = self.doc.add_paragraph(f"LinkedIn: {linkedin}", style='Normal')
            linkedin_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        github = pi.get("github")
        if github:
            github_line = self.doc.add_paragraph(f"GitHub: {github}", style='Normal')
            github_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        visa_line = self.doc.add_paragraph(pi.get("visa_status", ""), style='Normal')
        visa_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()  # blank line

        # ---- 2. Work Experience ----
        self._add_heading_with_line("Work Experience")
        experience_list = self.config.get('experience', [])

        for exp in experience_list:
            subhead_para = self.doc.add_paragraph(style='CustomSubheading')
            p_format = subhead_para.paragraph_format
            p_format.tab_stops.add_tab_stop(Inches(6), WD_TAB_ALIGNMENT.RIGHT)

            company_text = f"{exp.get('company', '')}, {exp.get('location', '')}"
            subhead_para.add_run(company_text)
            subhead_para.add_run("\t")
            subhead_para.add_run(exp.get('dates', ''))

            title_para = self.doc.add_paragraph(style='Normal')
            title_para.add_run(exp.get('title', '')).bold = True

            bullets = exp.get('bullet_points', [])
            for b in bullets:
                self.add_bullet(b)

            self.doc.add_paragraph()

        # ---- 3. Education ----
        education_list = self.config.get('education', [])
        if education_list:
            self._add_heading_with_line("Education")
            for edu in education_list:
                p = self.doc.add_paragraph(style='CustomSubheading')
                p_format = p.paragraph_format
                p_format.tab_stops.add_tab_stop(Inches(6), WD_TAB_ALIGNMENT.RIGHT)
                p.add_run(f"{edu.get('institution', '')} | {edu.get('degree', '')}")
                p.add_run("\t")
                p.add_run(edu.get('dates', ''))

                cgpa = edu.get('cgpa')
                if cgpa:
                    self.doc.add_paragraph(f"CGPA: {cgpa}", style='Normal')

            self.doc.add_paragraph()
        
        # ---- 4. Skills ----
        self._add_heading_with_line("Skills")

        categories = self.config.get('skills', {}).get('categories', [])
        for cat in categories:
            cat_name = cat.get('name', '')
            items = cat.get('items', [])
            cat_str = ", ".join(items)
            self._add_bold_label_value(f"{cat_name}: ", cat_str)

        if self.extra_skills:
            self.doc.add_paragraph("Additional Relevant Skills:")
            for skill in self.extra_skills:
                self.add_bullet(skill)

        spacer2 = self.doc.add_paragraph()
        spacer2.paragraph_format.space_before = Pt(0)
        spacer2.paragraph_format.space_after = Pt(2)
        spacer2.paragraph_format.line_spacing = 0.5

        # ---- 5. Certifications ----
        certs = self.config.get('certifications', [])
        if certs:
            self._add_heading_with_line("Certifications")
            for c in certs:
                self.add_bullet(c)

            spacer = self.doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(0)
            spacer.paragraph_format.space_after = Pt(2)
            spacer.paragraph_format.line_spacing = 0.5

        # ---- 6. Extracurricular Activities ----
        extras = self.config.get('extracurriculars', [])
        if extras:
            self._add_heading_with_line("Extracurricular Activities")
            for activity in extras:
                p = self.doc.add_paragraph(style='CustomSubheading')
                p_format = p.paragraph_format
                p_format.tab_stops.add_tab_stop(Inches(6), WD_TAB_ALIGNMENT.RIGHT)
                p.add_run(f"{activity.get('organization', '')} | {activity.get('position', '')}")
                p.add_run("\t")
                p.add_run(activity.get('dates', ''))

                for bullet in activity.get('bullet_points', []):
                    self.add_bullet(bullet)

            self.doc.add_paragraph()

        # ---- 7. Projects ----
        projects = self.config.get('projects', [])
        if projects:
            self._add_heading_with_line("Projects")
            for project in projects:
                p = self.doc.add_paragraph(style='CustomSubheading')
                p.add_run(project.get('name', ''))
                tech = project.get('tech_stack', [])
                if tech:
                    self.doc.add_paragraph(f"Tech Stack: {', '.join(tech)}", style='Normal')
                for bullet in project.get('bullet_points', []):
                    self.add_bullet(bullet)

            self.doc.add_paragraph()

    def save(self):
        # Save DOCX first
        self.doc.save(self.output_filename)
        print(f"{self.output_filename} has been generated successfully!")
    
        # Convert to PDF
        if self.output_filename.endswith(".docx"):
            pdf_filename = self.output_filename.replace(".docx", ".pdf")
            convert(self.output_filename, pdf_filename)
            print(f"{pdf_filename} has been generated successfully!")


def generate_mckinsey_style_cv(config=None, extra_skills=None):
    """
    Convenience function preserving the old signature.
    - config: dictionary loaded from YAML
    - extra_skills: list of new skills discovered from OpenAI
    """
    if config is None:
        raise ValueError("generate_mckinsey_style_cv(...) requires a config dictionary!")
    cv_gen = McKinseyCVGenerator(config=config, extra_skills=extra_skills)
    cv_gen.build()
    cv_gen.save()


if __name__ == "__main__":
    # If someone runs generate_cv.py directly (without passing config),
    # you can do a fallback or just raise an error:
    print("Please run this module via file1.py or import and pass a config dictionary.")
